"""Generate a user guideline document (DOCX) for the Journal Voucher System."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def generate_guide(output_path):
    doc = Document()

    # -- Styles --
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # =========================================================
    # COVER PAGE
    # =========================================================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Journal Voucher System')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('User Guide')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x37, 0x74, 0x9E)

    doc.add_paragraph()
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run('Version 1.0')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run('July 2026')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_page_break()

    # =========================================================
    # TABLE OF CONTENTS (manual)
    # =========================================================
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1.  Introduction',
        '2.  System Overview',
        '3.  Getting Started',
        '    3.1  Accessing the Application',
        '    3.2  User Login',
        '4.  Admin Dashboard (Jazzmin)',
        '    4.1  Navigation',
        '    4.2  Managing Master Data',
        '    4.3  Managing Accounting Periods',
        '    4.4  Managing Chart of Accounts',
        '    4.5  Managing Vendors',
        '    4.6  Managing Journal Vouchers',
        '5.  REST API Reference',
        '    5.1  Authentication',
        '    5.2  Master Data APIs',
        '    5.3  Accounting Period APIs',
        '    5.4  Chart of Account APIs',
        '    5.5  Vendor APIs',
        '    5.6  Journal Voucher APIs',
        '    5.7  Excel Upload API',
        '6.  Workflow Guide',
        '    6.1  Creating a Journal Voucher',
        '    6.2  Validating a Voucher',
        '    6.3  Voiding a Voucher',
        '    6.4  Batch Upload from Excel',
        '7.  Business Rules',
        '8.  Troubleshooting',
        '9.  Support',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = Pt(11)

    doc.add_page_break()

    # =========================================================
    # 1. INTRODUCTION
    # =========================================================
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'The Journal Voucher System is a web-based accounting application built with Django that '
        'supports journal voucher creation, validation, and management. It provides both an '
        'administrative dashboard for day-to-day operations and a REST API for programmatic access.'
    )
    doc.add_paragraph(
        'This guide provides step-by-step instructions for using the system, including navigating '
        'the admin interface, managing master data, working with journal vouchers, and consuming '
        'the REST API.'
    )

    # =========================================================
    # 2. SYSTEM OVERVIEW
    # =========================================================
    doc.add_heading('2. System Overview', level=1)
    doc.add_paragraph('The system is organized into the following functional modules:')

    add_styled_table(doc,
        ['Module', 'Purpose'],
        [
            ['Master Data', 'Companies, branches, and currencies used across the system'],
            ['Accounting Periods', 'Define open/closed periods for financial transactions'],
            ['Chart of Accounts', 'Hierarchical account structure (Assets, Liabilities, Equity, Revenue, Expense)'],
            ['Vendors', 'Vendor/supplier information including tax IDs, bank accounts, and credit limits'],
            ['Journal Vouchers', 'Core voucher engine with debit/credit entries, validation, and Excel upload'],
        ],
        col_widths=[4.5, 12],
    )

    doc.add_paragraph()
    doc.add_paragraph('Technology stack:')
    add_styled_table(doc,
        ['Layer', 'Technology'],
        [
            ['Backend', 'Django 5.2, Django REST Framework 3.15'],
            ['Admin UI', 'Jazzmin 3.0 (themed admin interface)'],
            ['Database', 'SQLite (development) / PostgreSQL (production)'],
            ['WSGI Server', 'Gunicorn 23'],
            ['Static Files', 'WhiteNoise 6'],
            ['Excel Processing', 'openpyxl'],
        ],
        col_widths=[4.5, 12],
    )

    # =========================================================
    # 3. GETTING STARTED
    # =========================================================
    doc.add_heading('3. Getting Started', level=1)

    doc.add_heading('3.1  Accessing the Application', level=2)
    doc.add_paragraph('The application is accessible via a web browser at the configured URL:')
    p = doc.add_paragraph()
    run = p.add_run('Admin Dashboard:  https://<hostname>/admin/')
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    p = doc.add_paragraph()
    run = p.add_run('API Endpoint:    https://<hostname>/api/')
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_heading('3.2  User Login', level=2)
    doc.add_paragraph('Users must authenticate to access the system. Login credentials are provided by the system administrator.')
    doc.add_paragraph('To log in:')
    steps = [
        'Navigate to the admin dashboard URL (/admin/).',
        'Enter your username and password.',
        'Click "Log in".',
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number')

    doc.add_paragraph(
        'For API access, authentication is handled via session authentication (browser) or '
        'basic authentication (programmatic clients). See Section 5.1 for details.'
    )

    # =========================================================
    # 4. ADMIN DASHBOARD
    # =========================================================
    doc.add_heading('4. Admin Dashboard (Jazzmin)', level=1)

    doc.add_heading('4.1  Navigation', level=2)
    doc.add_paragraph(
        'The admin interface uses the Jazzmin theme with a dark sidebar navigation panel. '
        'The sidebar contains links to all registered models, organized by application:'
    )
    nav_items = [
        'Master Data — Companies, Branches, Currencies',
        'Accounting Periods — Period lifecycle management',
        'Chart of Accounts — Account hierarchy',
        'Vendors — Vendor management',
        'Journal Vouchers — Core voucher engine',
        'Users & Groups — Django authentication',
    ]
    for item in nav_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('4.2  Managing Master Data', level=2)
    doc.add_paragraph('Master Data includes three models: Company, Branch, and Currency.')

    doc.add_heading('Companies', level=3)
    doc.add_paragraph('To add a new company:')
    steps = [
        'Navigate to Master Data > Companies in the sidebar.',
        'Click "Add Company" (top-right).',
        'Enter the Company Code (unique identifier), Name, Address, Phone, and Email.',
        'Set Active status (defaults to Active).',
        'Click "Save".',
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number')

    doc.add_heading('Branches', level=3)
    doc.add_paragraph(
        'Branches are linked to a parent Company. The combination of Company + Branch Code '
        'must be unique. Follow the same process as Companies to add branches.'
    )

    doc.add_heading('Currencies', level=3)
    doc.add_paragraph(
        'Currencies define the monetary units used in the system. Each currency has a unique '
        'code (e.g., USD, EUR), name, and optional symbol.'
    )

    doc.add_heading('4.3  Managing Accounting Periods', level=2)
    doc.add_paragraph(
        'Accounting periods define the time boundaries for financial transactions. '
        'Periods can be OPEN (accepting vouchers) or CLOSED (locked).'
    )
    doc.add_paragraph('To create a period:')
    steps = [
        'Navigate to Accounting Periods > Add.',
        'Enter a unique Period Code and Name.',
        'Set Start Date and End Date (start must be on or before end).',
        'Add optional notes.',
        'Click "Save". The period defaults to OPEN status.',
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number')

    doc.add_paragraph(
        'IMPORTANT: Vouchers can only be created in OPEN periods. Close a period after all '
        'transactions for that period are finalized.'
    )

    doc.add_heading('4.4  Managing Chart of Accounts', level=2)
    doc.add_paragraph(
        'The Chart of Accounts defines the hierarchical account structure. Accounts can have '
        'a parent account, enabling a tree-like structure (e.g., Assets > Current Assets > Cash).'
    )
    doc.add_paragraph('Account types:')
    add_styled_table(doc,
        ['Type', 'Normal Balance', 'Description'],
        [
            ['ASSET', 'Debit', 'Resources owned by the entity'],
            ['LIABILITY', 'Credit', 'Obligations owed to others'],
            ['EQUITY', 'Credit', 'Owner\'s interest in the entity'],
            ['REVENUE', 'Credit', 'Income from operations'],
            ['EXPENSE', 'Debit', 'Costs incurred in operations'],
        ],
        col_widths=[3.5, 3.5, 9.5],
    )

    doc.add_heading('4.5  Managing Vendors', level=2)
    doc.add_paragraph(
        'Vendors store supplier information. Fields include code, name, contact details, '
        'tax ID, bank account, credit limit, and payment terms.'
    )

    doc.add_heading('4.6  Managing Journal Vouchers', level=2)
    doc.add_paragraph(
        'The Journal Voucher admin page provides a comprehensive view of all vouchers with '
        'inline entry management.'
    )
    doc.add_paragraph('Features:')
    features = [
        'List view — filterable by status and period, searchable by voucher number and description.',
        'Inline entries — add or edit journal entries directly within the voucher form.',
        'Auto-generated voucher numbers in the format JV-YYYYMMDD-NNNN.',
        'Automatic calculation of total debit and credit from entries.',
        'Date hierarchy navigation for browsing by transaction date.',
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')

    # =========================================================
    # 5. REST API REFERENCE
    # =========================================================
    doc.add_heading('5. REST API Reference', level=1)
    doc.add_paragraph(
        'The system exposes a comprehensive REST API built with Django REST Framework. '
        'All endpoints return JSON responses and support pagination (25 items per page), '
        'searching (?search=), and ordering (?ordering=).'
    )

    doc.add_heading('5.1  Authentication', level=2)
    doc.add_paragraph('The API supports two authentication methods:')
    methods = [
        'Session Authentication — used automatically when logged in via the admin dashboard.',
        'Basic Authentication — for programmatic clients. Include an Authorization header.',
    ]
    for m in methods:
        doc.add_paragraph(m, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('All API endpoints require authentication.')
    run.bold = True

    doc.add_heading('5.2  Master Data APIs', level=2)
    add_styled_table(doc,
        ['Method', 'Endpoint', 'Description'],
        [
            ['GET, POST', '/api/master-data/companies/', 'List / Create companies'],
            ['GET, PUT, PATCH, DELETE', '/api/master-data/companies/{id}/', 'Company detail'],
            ['GET, POST', '/api/master-data/branches/', 'List / Create branches'],
            ['GET, PUT, PATCH, DELETE', '/api/master-data/branches/{id}/', 'Branch detail'],
            ['GET, POST', '/api/master-data/currencies/', 'List / Create currencies'],
            ['GET, PUT, PATCH, DELETE', '/api/master-data/currencies/{id}/', 'Currency detail'],
        ],
        col_widths=[3.5, 6, 7],
    )

    doc.add_heading('5.3  Accounting Period APIs', level=2)
    add_styled_table(doc,
        ['Method', 'Endpoint', 'Description'],
        [
            ['GET, POST', '/api/accounting-periods/', 'List / Create periods'],
            ['GET, PUT, PATCH, DELETE', '/api/accounting-periods/{id}/', 'Period detail'],
        ],
        col_widths=[3.5, 6, 7],
    )

    doc.add_heading('5.4  Chart of Account APIs', level=2)
    add_styled_table(doc,
        ['Method', 'Endpoint', 'Description'],
        [
            ['GET, POST', '/api/chart-of-accounts/', 'List / Create accounts'],
            ['GET, PUT, PATCH, DELETE', '/api/chart-of-accounts/{id}/', 'Account detail'],
        ],
        col_widths=[3.5, 6, 7],
    )

    doc.add_heading('5.5  Vendor APIs', level=2)
    add_styled_table(doc,
        ['Method', 'Endpoint', 'Description'],
        [
            ['GET, POST', '/api/vendors/', 'List / Create vendors'],
            ['GET, PUT, PATCH, DELETE', '/api/vendors/{id}/', 'Vendor detail'],
        ],
        col_widths=[3.5, 6, 7],
    )

    doc.add_heading('5.6  Journal Voucher APIs', level=2)
    add_styled_table(doc,
        ['Method', 'Endpoint', 'Description'],
        [
            ['GET', '/api/journal-vouchers/', 'List vouchers (paginated)'],
            ['POST', '/api/journal-vouchers/', 'Create voucher with nested entries'],
            ['GET', '/api/journal-vouchers/{id}/', 'Retrieve voucher with entries'],
            ['PUT, PATCH', '/api/journal-vouchers/{id}/', 'Update voucher / replace entries'],
            ['DELETE', '/api/journal-vouchers/{id}/', 'Delete voucher (entries cascade)'],
            ['POST', '/api/journal-vouchers/{id}/validate/', 'Validate DRAFT voucher'],
            ['POST', '/api/journal-vouchers/{id}/void/', 'Void a voucher'],
            ['POST', '/api/journal-vouchers/upload_excel/', 'Batch create from .xlsx'],
        ],
        col_widths=[3.5, 6, 7],
    )

    doc.add_heading('5.7  Excel Upload API', level=2)
    doc.add_paragraph(
        'The upload_excel endpoint accepts a multipart POST with an .xlsx file. '
        'Expected columns:'
    )
    add_styled_table(doc,
        ['Column', 'Header', 'Required', 'Description'],
        [
            ['A', 'Voucher Description', 'Yes', 'Groups rows into one balanced voucher'],
            ['B', 'Account Code', 'Yes', 'Must match an existing ChartOfAccount.code'],
            ['C', 'Debit', 'Conditional', 'Debit amount (0 if credit entry)'],
            ['D', 'Credit', 'Conditional', 'Credit amount (0 if debit entry)'],
            ['E', 'Line Description', 'No', 'Line-item description'],
            ['F', 'Transaction Date', 'No', 'Defaults to today if omitted'],
            ['G', 'Accounting Period Code', 'Yes', 'Must match an existing AccountingPeriod.code'],
        ],
        col_widths=[2, 4, 2, 8.5],
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Request format:')
    run.bold = True
    code = (
        'POST /api/journal-vouchers/upload_excel/\n'
        'Content-Type: multipart/form-data\n\n'
        'file: @vouchers.xlsx'
    )
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Response:')
    run.bold = True
    code = (
        '{\n'
        '  "created": 3,\n'
        '  "errors": [\n'
        '    "Voucher X: error description"\n'
        '  ]\n'
        '}'
    )
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    # =========================================================
    # 6. WORKFLOW GUIDE
    # =========================================================
    doc.add_heading('6. Workflow Guide', level=1)

    doc.add_heading('6.1  Creating a Journal Voucher', level=2)
    doc.add_paragraph('Using the REST API:')
    code = (
        'POST /api/journal-vouchers/\n'
        'Content-Type: application/json\n\n'
        '{\n'
        '  "accounting_period": 1,\n'
        '  "transaction_date": "2026-01-15",\n'
        '  "description": "January revenue entry",\n'
        '  "vendor": 1,\n'
        '  "entries": [\n'
        '    {"account": 1, "debit": 1000.00, "credit": 0, "description": "Cash debit"},\n'
        '    {"account": 2, "debit": 0, "credit": 1000.00, "description": "Revenue credit"}\n'
        '  ]\n'
        '}'
    )
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_paragraph()
    doc.add_paragraph('Rules for entries:')
    rules = [
        'Each entry must have either debit > 0 or credit > 0 (not both, not neither).',
        'Sum of all debits must equal sum of all credits across entries.',
        'The account must exist and be active.',
        'The accounting period must exist and be OPEN.',
        'The transaction date must fall within the accounting period\'s date range.',
    ]
    for r in rules:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_heading('6.2  Validating a Voucher', level=2)
    doc.add_paragraph(
        'A DRAFT voucher can be validated to confirm it meets all business rules. '
        'Validation checks:'
    )
    checks = [
        'Voucher status is DRAFT (cannot validate already validated, posted, or voided vouchers).',
        'Voucher has at least one journal entry.',
        'Accounting period is OPEN.',
        'Total debit equals total credit.',
        'All accounts used in entries are ACTIVE.',
    ]
    for c in checks:
        doc.add_paragraph(c, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('POST /api/journal-vouchers/{id}/validate/')
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_paragraph(
        'On success, the voucher status changes to VALIDATED and the updated voucher is returned. '
        'On failure, a 400 response with error details is returned.'
    )

    doc.add_heading('6.3  Voiding a Voucher', level=2)
    doc.add_paragraph(
        'A voucher can be voided at any point in its lifecycle (DRAFT, VALIDATED, or POSTED). '
        'Once voided, the status becomes VOID and no further actions can be performed.'
    )
    p = doc.add_paragraph()
    run = p.add_run('POST /api/journal-vouchers/{id}/void/')
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_heading('6.4  Batch Upload from Excel', level=2)
    doc.add_paragraph(
        'The Excel upload feature allows creating multiple vouchers in a single batch. '
        'Rows with the same "Voucher Description" are grouped into one voucher with multiple entries.'
    )
    doc.add_paragraph('Steps:')
    steps = [
        'Prepare an .xlsx file with columns A through G as shown in Section 5.7.',
        'Send a POST request with the file as a multipart upload.',
        'The system validates each voucher group: debit = credit, accounts exist, period exists.',
        'On success, the response shows the number of created vouchers and any errors.',
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('Example Excel data:')
    add_styled_table(doc,
        ['Description', 'Account Code', 'Debit', 'Credit', 'Description', 'Date', 'Period Code'],
        [
            ['January Rent', '5100', '5000', '', 'Rent expense', '2026-01-01', '202601'],
            ['January Rent', '1000', '', '5000', 'Cash payment', '2026-01-01', '202601'],
            ['Consulting Fees', '4000', '', '2000', 'Consulting revenue', '2026-01-15', '202601'],
            ['Consulting Fees', '1000', '2000', '', 'Bank deposit', '2026-01-15', '202601'],
        ],
        col_widths=[3, 2.5, 1.5, 1.5, 3, 2.5, 2.5],
    )

    # =========================================================
    # 7. BUSINESS RULES
    # =========================================================
    doc.add_heading('7. Business Rules', level=1)
    rules = [
        ('Debit = Credit', 'Every journal voucher must have total debit equal to total credit. This is the fundamental accounting equation for double-entry bookkeeping.'),
        ('Period must be OPEN', 'Vouchers can only be created and validated in OPEN periods. CLOSED periods are locked and prevent modifications.'),
        ('Account must be ACTIVE', 'Only ACTIVE accounts can be used in journal entries. INACTIVE accounts are preserved for historical data but cannot be used in new transactions.'),
        ('Transaction date in range', 'The transaction date must fall within the linked accounting period\'s start and end dates.'),
        ('Unique voucher numbers', 'Voucher numbers are auto-generated in the format JV-YYYYMMDD-NNNN and are guaranteed to be unique.'),
        ('Voucher lifecycle', 'A voucher progresses through statuses: DRAFT > VALIDATED > POSTED > VOID. Validation and voiding are irreversible actions.'),
    ]
    for title, desc in rules:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        p.add_run(desc)

    # =========================================================
    # 8. TROUBLESHOOTING
    # =========================================================
    doc.add_heading('8. Troubleshooting', level=1)
    add_styled_table(doc,
        ['Issue', 'Likely Cause', 'Solution'],
        [
            ['Cannot log in', 'Invalid credentials or inactive account',
             'Reset password via admin or contact system administrator.'],
            ['Cannot create voucher', 'Period is CLOSED or transaction date out of range',
             'Verify the period is OPEN and the date falls within period boundaries.'],
            ['Validation fails: debit != credit', 'Journal entries are imbalanced',
             'Check that total debit equals total credit across all entries.'],
            ['Validation fails: account not active', 'An entry references an INACTIVE account',
             'Use only ACTIVE accounts, or reactivate the account via admin.'],
            ['Excel upload creates 0 vouchers', 'Column format or data mismatch',
             'Verify column order matches the expected format. Check response errors array.'],
            ['API returns 403', 'Missing or invalid authentication',
             'Ensure you are logged in or include valid Basic Auth headers.'],
            ['API returns 404', 'Invalid endpoint or resource ID',
             'Verify the URL and resource ID. Check the endpoint list in Section 5.'],
        ],
        col_widths=[4, 4.5, 8],
    )

    # =========================================================
    # 9. SUPPORT
    # =========================================================
    doc.add_heading('9. Support', level=1)
    doc.add_paragraph(
        'For technical support, bug reports, or feature requests, please contact the system '
        'administrator or development team.'
    )
    doc.add_paragraph('When reporting an issue, please include:')
    items = [
        'The exact URL and HTTP method used.',
        'The request payload (if applicable).',
        'The full error response (including HTTP status code and response body).',
        'Steps to reproduce the issue.',
        'The browser or client being used.',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # =========================================================
    # SAVE
    # =========================================================
    doc.save(output_path)
    print(f'User guide generated: {output_path}')


if __name__ == '__main__':
    output = os.path.join(os.path.dirname(__file__), 'Journal_Voucher_System_User_Guide.docx')
    generate_guide(output)
